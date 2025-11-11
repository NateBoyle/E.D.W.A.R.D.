#!/usr/bin/env python
# coding: utf-8

# In[63]:


# import os (operating system) library
import os

# import sys (system) library
import sys

# import pandas library as pd
import pandas as pd


# import numpy library as np
import numpy as np

# import these modules from datetime to store and compare dates
from datetime import datetime, date, timedelta

# import time for delay
import time
import pygame

# from tkinter import all standard modules with * (this is for the gui)
from tkinter import *
# from tkinter 'specifically' import messagebox, ttk as they are not standard modules uploaded with *
from tkinter import messagebox, ttk
from tkinter import filedialog as fd

# import imageTk and Image for picture use in gui
from PIL import ImageTk as itk, Image 

# import various matplotlib modules to create plots and then draw those plots in the gui
import matplotlib
import matplotlib.pyplot as plt
from matplotlib.figure import Figure 
import matplotlib.ticker as mtick
from matplotlib.backends.backend_tkagg import (FigureCanvasTkAgg, NavigationToolbar2Tk)

from docx import Document
from docx.shared import Inches, Cm
from docx.shared import Pt
from docxcompose.composer import Composer

from docxtpl import DocxTemplate
import jinja2

import random
import ipyplot

# import webbrowser for hyperlink use
import webbrowser

from docx2pdf import convert
from tkPDFViewer import tkPDFViewer as pdf 

from io import StringIO
from io import BytesIO


# import IPython display for wider coding screen (not required to run program)
from IPython.display import display, HTML
display(HTML("<style>.jp-Cell { width: 120% !important; }</style>"))


# In[64]:


def plotAccountGraph(df, acclist):

    plotdf = df.iloc[:,2:].set_index('acct_name')

    fig, ax = plt.subplots(figsize=(14, 8))

    
    plot1df = plotdf.loc[acclist[0],:]

    ax.plot(plot1df)

    ax.get_yaxis().set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x, p: format(int(x/1000), ',')))

    
    ax.get_yaxis().set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x, p: format(int(x/1000), ',')))
    ax.set_ylabel("USD in 000s", fontsize="12")
    ax.set_xlabel('Year', fontsize="12")

    
    if len(acclist) == 2:
        plot2df = plotdf.loc[acclist[1],:]
        ax.plot(plot2df)

    if len(acclist) == 3:
        plot2df = plotdf.loc[acclist[1],:]
        ax.plot(plot2df)
        plot3df = plotdf.loc[acclist[2],:]
        ax.plot(plot3df)

    if len(acclist) == 1:
        titlestr = acclist[0] + " change over time."
    elif len(acclist) == 2:
        titlestr = acclist[0] + " and " + acclist[1] + " change over time."
    elif len(acclist) == 3:
        titlestr = acclist[0] + ", " + acclist[1] + ", and " + acclist[2] + " change over time."

    ax.set_title(titlestr)

    ax.legend(acclist, ncol=len(acclist), loc="upper center", bbox_to_anchor=(0.5, -0.1),
          fancybox=True, fontsize="12")
    
    memfile = BytesIO()
    
    
    plt.savefig(memfile, bbox_inches="tight")

    plt.close(fig)


    return memfile

    
    


# In[65]:


def plotAccountCharts(df1,df2,df3,str1,str2,str3,vh,n):

    if n == 2:
        fig, axes = plt.subplots(1,2, figsize=(14, 6))
    if n == 3:
        fig, axes = plt.subplots(1,3, figsize=(14, 6))
        axes[2].set_title(str3)
        

    if vh == 'v':
        
        df1.plot.bar(ax=axes[0], stacked=True)
        axes[0].set_ylabel("USD in 000s", fontsize="12")
        
        df2.plot.bar(ax=axes[1], stacked=True)
 
        for i in range(2):
            axes[i].get_yaxis().set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x, p: format(int(x/1000), ',')))
            #axes[i].set_xlabel('Year', fontsize="12")
            axes[i].legend(ncol=2, loc="upper center", bbox_to_anchor=(0.5, -0.15),
              fancybox=True, fontsize="9")
        
        if n == 3:
            df3.plot.bar(ax=axes[2], stacked=True)
            axes[2].get_yaxis().set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x, p: format(int(x/1000), ',')))
            #axes[2].set_xlabel('Year', fontsize="12")
            axes[2].legend(ncol=2, loc="upper center", bbox_to_anchor=(0.5, -0.1),
              fancybox=True, fontsize="9")
            
        
    elif vh == 'h':

        df1.plot.barh(ax=axes[0], stacked=True)
        axes[0].set_ylabel("Year", fontsize="12")
        
        df2.plot.barh(ax=axes[1], stacked=True)
 
        for i in range(2):
            axes[i].set_xlabel('Percent', fontsize="12")
            axes[i].legend(ncol=2, loc="upper center", bbox_to_anchor=(0.5, -0.15),
              fancybox=True, fontsize="10")
        
        if n == 3:
            df3.plot.barh(ax=axes[2], stacked=True)
            axes[2].set_xlabel('Percent', fontsize="12")
            axes[2].legend(ncol=2, loc="upper center", bbox_to_anchor=(0.5, -0.15),
              fancybox=True, fontsize="10")
            

    axes[0].set_title(str1)
    axes[1].set_title(str2)
        
    
    memfile = BytesIO()
    
    # memlist = [' ']*2
    # memlist[0] = liqmemfile
    plt.savefig(memfile, bbox_inches="tight")

    plt.close(fig)

    return memfile



# In[66]:


def plotCF(df, years):

    cfdict = df.to_dict('list')

    # cash flow composition graph
    cflist = [0]*len(cfdict)
    counter = 0
    
    for key in cfdict:
        cflist[counter] = cfdict[key]
        counter += 1
    
    data = np.array(cflist)
    
    data_shape = np.shape(data)
    
    # Take negative and positive data apart and cumulate
    def get_cumulated_array(data, **kwargs):
        cum = data.clip(**kwargs)
        cum = np.cumsum(cum, axis=0)
        d = np.zeros(np.shape(data))
        d[1:] = cum[:-1]
        return d
    
    cumulated_data = get_cumulated_array(data, min=0)
    cumulated_data_neg = get_cumulated_array(data, max=0)
    
    # Re-merge negative and positive data.
    row_mask = (data<0)
    cumulated_data[row_mask] = cumulated_data_neg[row_mask]
    data_stack = cumulated_data
    
    width = 0.5
    
    fig, ax = plt.subplots(1,1, figsize=(14, 7))
    
    
    cfkeylist = list(cfdict.keys())
    
    for i in np.arange(0, data_shape[0]):
        ax.bar(years, data[i], bottom=data_stack[i], label=cfkeylist[i])
    
    # Shrink current axis's height by 10% on the bottom
    box = ax.get_position()
    ax.set_position([box.x0, box.y0 + box.height * 0.1,
                     box.width, box.height * 0.9])
    
    ax.get_yaxis().set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x, p: format(int(x/1000), ',')))
    ax.set_ylabel("USD in 000s", fontsize="12")
    ax.set_xlabel('Year', fontsize="12")
    
    ax.set_title("Cash Flow Composition")
    ax.legend(ncol=2, loc="upper center", bbox_to_anchor=(0.5, -0.1),
              fancybox=True, fontsize="10")
    
    
    
    #plt.show()

    memfile = BytesIO()
    
    # memlist = [' ']*2
    # memlist[0] = liqmemfile
    plt.savefig(memfile, bbox_inches="tight")

    plt.close(fig)

    return memfile
        


# In[67]:


def createCharts(ac,fr,at,df,acclist):

    memlist = [0]*10
    assetfile = ''
    liabfile = ''
    aleisfile = ''
    ilefile = ''
    cffile = ''
    pmfile = ''
    liqfile = ''
    solvfile = ''
    trendfile = ''
    

    tbdf = df
    years = tbdf.columns[3:].tolist()

    if ac:
        # assets composition
        cadf = tbdf[(tbdf['acct_key'] == 'CURRENT_ASSET')].iloc[:,2:].set_index('acct_name').T
        ncadf = tbdf[(tbdf['acct_key'] == 'NONCURRENT_ASSET')].iloc[:,2:].set_index('acct_name').T
    
        str1 = 'Current Assets Composition'
        str2 = 'Noncurrent Assets Composition'
        
        assetfile = plotAccountCharts(cadf,ncadf,None,str1,str2,None,'v',2)
    
    
        # liabilities composition
        cldf = tbdf[(tbdf['acct_key'] == 'CURRENT_LIABILITY')].iloc[:,2:].set_index('acct_name').T
        ncldf = tbdf[(tbdf['acct_key'] == 'NONCURRENT_LIABILITY')].iloc[:,2:].set_index('acct_name').T
    
        str1 = 'Current Liabilities Composition'
        str2 = 'Noncurrent Liabilities Composition'
        
        liabfile = plotAccountCharts(cldf,ncldf,None,str1,str2,None, 'v',2)
    
        
        # balance sheet and income statement composition
        alelist = ['TOT_ASSET', 'TOT_LIABILITY', 'EQUITY']
        aledf = tbdf[tbdf['acct_key'].isin(alelist)].iloc[:,2:].set_index('acct_name').T
        
        islist = ['TOT_REVENUE', 'TOT_COS', 'TOT_OPEX']
        isdf = tbdf[tbdf['acct_key'].isin(islist)].iloc[:,2:].set_index('acct_name').T
    
        str1 = 'Balance Sheet Composition'
        str2 = 'Income Statement Composition'
        
        aleisfile = plotAccountCharts(aledf,isdf,None,str1,str2,None, 'v',2)
    
        
        # income and loss and expense composition
        revdf = tbdf[tbdf['acct_key'] == 'REVENUE'].iloc[:,2:]
        
        oiedf = tbdf[(tbdf['acct_key'] == 'OTHER_INCOME_EXPENSE') ].iloc[:,2:]
        oidf = oiedf.copy()
        oidf[oidf[years] < 0] = 0
        
        explist = ['COS', 'OPEX']
        expdf = tbdf[tbdf['acct_key'].isin(explist)].iloc[:,2:]
        
        gaindf = expdf.copy()
        
        expdf[expdf[years] < 0] = 0
        gaindf[gaindf[years] > 0] = 0
        
        gaindf[gaindf.select_dtypes(include=['number']).columns] = gaindf[gaindf.select_dtypes(include=['number']).columns].abs()
        
        oedf = oiedf.copy()
        oedf[oedf[years] > 0] = 0
        oedf[oedf.select_dtypes(include=['number']).columns] = oedf[oedf.select_dtypes(include=['number']).columns].abs()
        
        revdf = pd.concat([revdf, gaindf, oidf])
        revdf = revdf[(revdf.iloc[:,2:].T != 0).any()].set_index('acct_name').T
        
        expdf = pd.concat([expdf, oedf])
        expdf = expdf[(expdf.iloc[:,2:].T != 0).any()].set_index('acct_name').T
    
        str1 = 'Income Composition'
        str2 = 'Loss and Expense Composition'
        
        ilefile = plotAccountCharts(revdf,expdf,None,str1,str2,None, 'v',2)
        
        
        cfdf = tbdf[tbdf['acct_key'] == 'CF'].iloc[:,2:].set_index('acct_name').T
    
        cffile = plotCF(cfdf, years)

    if fr:

        ## profit margin ratios
        pmlist = ['TOT_REVENUE', 'TOT_COS', 'TOT_OPEX']
        pmdf = tbdf[tbdf['acct_key'].isin(pmlist)].iloc[:,2:]
        # operating margin
        opmargindf = pmdf.set_index('acct_name')
        # gross margin
        gpmargindf = pmdf.set_index('acct_name').loc[['Total Revenue','Total Cost of Revenues'],:]
        
        #Total sum per row: 
        opmargindf.loc['Total',:] = opmargindf.sum(axis=0)
        gpmargindf.loc['Total',:] = gpmargindf.sum(axis=0)
        
        # create new rows for each original row as a percent of the new total row
        opmargindf.loc['Revenue Percent',:] = opmargindf.loc['Total Revenue',:]/opmargindf.loc['Total',:]*100
        opmargindf.loc['COS Percent',:] = opmargindf.loc['Total Cost of Revenues',:]/opmargindf.loc['Total',:]*100
        opmargindf.loc['OPEX Percent',:] = opmargindf.loc['Total Operating Expenses',:]/opmargindf.loc['Total',:]*100
        
        gpmargindf.loc['Revenue Percent',:] = gpmargindf.loc['Total Revenue',:]/gpmargindf.loc['Total',:]*100
        gpmargindf.loc['COS Percent',:] = gpmargindf.loc['Total Cost of Revenues',:]/gpmargindf.loc['Total',:]*100
        
        # just keep percent rows
        opmargindf = opmargindf.loc[['Revenue Percent','COS Percent','OPEX Percent'],:].T

        gpmargindf = gpmargindf.loc[['Revenue Percent','COS Percent'],:].T

        str1 = 'Operating Margin Ratio'
        str2 = 'Gross Margin Ratio'

        pmfile = plotAccountCharts(opmargindf,gpmargindf,None,str1,str2,None, 'h',2)


        ## Liquidity ratios

        # current ratio
        currentdf = tbdf[(tbdf['acct_key'] == 'TOT_CURRENT_ASSET') | (tbdf['acct_key'] == 'TOT_CURRENT_LIABILITY')].iloc[:,2:].set_index('acct_name')
        currentdf.loc['Total',:] = currentdf.sum(axis=0)
        currentdf = currentdf.div(currentdf.loc['Total',:])*100
        currentdf = currentdf.loc[~currentdf.index.isin(['Total'])].T

        # quick ratio
        quickdf = tbdf[(tbdf['acct_key'] == 'CURRENT_ASSET') | (tbdf['acct_key'] == 'TOT_CURRENT_LIABILITY')].iloc[:,2:].set_index('acct_name')
        quickdf = quickdf.loc[~quickdf.index.isin(['Inventory','Prepaid Expenses and Other Current Assets'])]
        quickdf.loc['Total',:] = quickdf.sum(axis=0)
        quickdf = quickdf.div(quickdf.loc['Total',:])*100
        quickdf = quickdf.loc[~quickdf.index.isin(['Total'])].T

        str1 = 'Current Ratio'
        str2 = 'Quick Ratio'

        liqfile = plotAccountCharts(currentdf,quickdf,None,str1,str2,None, 'h',2)

        
        ## Solvency ratios
        
        aedf = tbdf[(tbdf['acct_key'] == 'TOT_ASSET') | (tbdf['acct_key'] == 'EQUITY')].iloc[:,2:].set_index('acct_name')
        aedf.loc['Total',:] = aedf.sum(axis=0)
        aedf = aedf.div(aedf.loc['Total',:])*100
        aedf = aedf.loc[~aedf.index.isin(['Total'])].T
        
        aldf = tbdf[(tbdf['acct_key'] == 'TOT_ASSET') | (tbdf['acct_key'] == 'TOT_LIABILITY')].iloc[:,2:].set_index('acct_name')
        aldf.loc['Total',:] = aldf.sum(axis=0)
        aldf = aldf.div(aldf.loc['Total',:])*100
        aldf = aldf.loc[~aldf.index.isin(['Total'])].T
        
        dedf = tbdf[(tbdf['acct_key'] == 'TOT_LIABILITY') | (tbdf['acct_key'] == 'EQUITY')].iloc[:,2:].set_index('acct_name')
        dedf.loc['Total',:] = dedf.sum(axis=0)
        dedf = dedf.div(dedf.loc['Total',:])*100
        dedf = dedf.loc[~dedf.index.isin(['Total'])].T
    
        str1 = 'Equity-Assets Ratio'
        str2 = 'Debt-Assets Ratio'
        str3 = 'Debt-Equity Ratio'
        
        solvfile = plotAccountCharts(aedf,aldf,dedf,str1,str2,str3,'h',3)


    if at:

        trendfile = plotAccountGraph(df, acclist)

    document= Document()
    sections = document.sections
    
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        
    if ac:
        
        document.add_heading("Account Composition Charts")
        
        p1 = document.add_paragraph("Asset Composition")
        p1.style = document.styles['Normal']
        r1 = p1.add_run()
        r1.add_picture(assetfile, width = Inches(7.5))
    
        p2 = document.add_paragraph("Liabilities Composition")
        p2.style = document.styles['Normal']
        r2 = p2.add_run()
        r2.add_picture(liabfile, width = Inches(7.5))
    
        p3 = document.add_paragraph("Balance Sheet and Income Statement Composition")
        p3.style = document.styles['Normal']
        r3 = p3.add_run()
        r3.add_picture(aleisfile, width = Inches(7.5))
    
        p4 = document.add_paragraph("Income and Loss and Expense Composition")
        p4.style = document.styles['Normal']
        r4 = p4.add_run()
        r4.add_picture(ilefile, width = Inches(7.5))
    
        p5 = document.add_paragraph("Cash Flow Composition")
        p5.style = document.styles['Normal']
        r5 = p5.add_run()
        r5.add_picture(cffile, width = Inches(7.5))

    if fr:

        if ac:
            document.add_page_break() 
        
        document.add_heading("Financial Ratio Visualizations")
        
        p6 = document.add_paragraph("Profit Margin Ratios")
        p6.style = document.styles['Normal']
        r6 = p6.add_run()
        r6.add_picture(pmfile, width = Inches(7.5))
    
        p7 = document.add_paragraph("Liquidity Ratios")
        p7.style = document.styles['Normal']
        r7 = p7.add_run()
        r7.add_picture(liqfile, width = Inches(7.5))
    
        p8 = document.add_paragraph("Solvency Ratios")
        p8.style = document.styles['Normal']
        r8 = p8.add_run()
        r8.add_picture(solvfile, width = Inches(7.5))

    if at:

        if ac or fr:
            document.add_page_break() 

        document.add_heading("Account Trend(s) Graph")
        
        p9 = document.add_paragraph("\n")
        p9.style = document.styles['Normal']
        r9 = p9.add_run()
        r9.add_picture(trendfile, width = Inches(7.5))
        
    
    #document.save("chartstest.docx")

    return document

    


# In[68]:


def createFS(df, name):

    #create financial statement df with just first and last two columns
    fsdf = df.iloc[:,list(range(0,3,1)) + list(range(-2,0,1))]
    cy = fsdf.iloc[:,list(range(-2,0,1))].columns.tolist()[1]
    py = fsdf.iloc[:,list(range(-2,0,1))].columns.tolist()[0]
    #format last two columns for currency
    fsdf.iloc[:,-1]=fsdf[cy].div(1000).apply('{:,.0f}'.format)
    fsdf.iloc[:,-2]=fsdf[py].div(1000).apply('{:,.0f}'.format)
    # keep only the accounts where both balances are nonzero
    fsdf = fsdf[~((fsdf[cy] == '0') & (fsdf[py] == '0'))]
    # replace 0 strings with '-'
    fsdf.loc[fsdf[cy] == '0', cy] = '-'
    fsdf.loc[fsdf[py] == '0', py] = '-'

    # create sub dictionaries for each multi fs line account group and add to fs dictionary
    fsdict = {}
    fsdict['CURRENT_ASSET'] = fsdf.query("acct_key == 'CURRENT_ASSET'").set_index('fs_key').T.to_dict('list') 
    fsdict['NONCURRENT_ASSET'] = fsdf.query("acct_key == 'NONCURRENT_ASSET'").set_index('fs_key').T.to_dict('list')
    fsdict['CURRENT_LIABILITY'] = fsdf.query("acct_key == 'CURRENT_LIABILITY'").set_index('fs_key').T.to_dict('list')
    fsdict['NONCURRENT_LIABILITY'] = fsdf.query("acct_key == 'NONCURRENT_LIABILITY'").set_index('fs_key').T.to_dict('list')
    fsdict['REVENUE'] = fsdf.query("acct_key == 'REVENUE'").set_index('fs_key').T.to_dict('list')
    fsdict['COS'] = fsdf.query("acct_key == 'COS'").set_index('fs_key').T.to_dict('list')
    fsdict['OPEX'] = fsdf.query("acct_key == 'OPEX'").set_index('fs_key').T.to_dict('list')
    fsdict['OTHER_INCOME_EXPENSE'] = fsdf.query("acct_key == 'OTHER_INCOME_EXPENSE'").set_index('fs_key').T.to_dict('list')
    fsdict['CF'] = fsdf.query("acct_key == 'CF'").set_index('fs_key').T.to_dict('list')

    # create dictionary for remaining single fs line items
    onefslinedict = fsdf[~fsdf["acct_key"].isin(fsdict.keys())].copy().set_index('acct_key').T.to_dict('list')

    if 'DISC_OPS' not in onefslinedict:
        onefslinedict['DISC_OPS'] = ['fsDO', 'Discontinued Operations', '-', '-']

    if 'INCOME_TAX' not in onefslinedict:
        onefslinedict['INCOME_TAX'] = ['fsITE', 'Income Tax Expense', '-', '-']
        

    # combine single and mutli fs line dictionaries
    fsdict = fsdict | onefslinedict

    # create additional fs line subtotals
    pyglfromop = float(fsdict['TOT_REVENUE'][2].replace(',','')) - float(fsdict['TOT_COS'][2].replace(',','')) - float(fsdict['TOT_OPEX'][2].replace(',',''))
    cyglfromop = float(fsdict['TOT_REVENUE'][3].replace(',','')) - float(fsdict['TOT_COS'][3].replace(',','')) - float(fsdict['TOT_OPEX'][3].replace(',',''))
    pyglb4tax = pyglfromop + float(fsdict['TOTAL_OTHER_INCOME_EXPENSE'][2].replace(',',''))
    cyglb4tax = cyglfromop + float(fsdict['TOTAL_OTHER_INCOME_EXPENSE'][3].replace(',',''))
    
    if fsdict['INCOME_TAX'][3].replace(',','').isnumeric():
        cyglfromcop = cyglb4tax - float(fsdict['INCOME_TAX'][3].replace(',',''))
    else:
        cyglfromcop = cyglb4tax
    
    if fsdict['INCOME_TAX'][2].replace(',','').isnumeric():
        pyglfromcop = pyglb4tax - float(fsdict['INCOME_TAX'][2].replace(',',''))
    else:
        pyglfromcop = pyglb4tax

    # add in additional fs line subtotals
    fsdict['CYGLFROMOP'] = '{:,.0f}'.format(cyglfromop)
    fsdict['PYGLFROMOP'] = '{:,.0f}'.format(pyglfromop)
    fsdict['CYGLB4TAX'] = '{:,.0f}'.format(cyglb4tax)
    fsdict['PYGLB4TAX'] = '{:,.0f}'.format(pyglb4tax)
    fsdict['CYGLFROMCOP'] = '{:,.0f}'.format(cyglfromcop)
    fsdict['PYGLFROMCOP'] = '{:,.0f}'.format(pyglfromcop)

    fsdict['cy'] = 'Fiscal Year '+cy
    fsdict['py'] = 'Fiscal Year '+py

    fsdict['name'] = name

    return fsdict
    


# In[73]:


def produceDocument(fs,ac,fr,at,df,name,acclist):

    f = ''
    fsdict = {}
    fsdoc = ''
    chartdoc = ''
    
    if fs:

        fsdict = createFS(df,name)
        
        fsdoc = DocxTemplate('fs_template.docx')
        fsdoc.render(fsdict)
        

    if any([ac,fr,at]):

        chartdoc = createCharts(ac,fr,at,df,acclist)

       

    f = fd.asksaveasfilename(defaultextension=".docx", title="Select name and path.")
    #print(f)
    if f:
        
        try:

            if fs and any([ac,fr,at]):
                
                composer = Composer(fsdoc)
                composer.append(chartdoc)
                composer.save(f)
                
            elif fs and not any([ac,fr,at]):
                
                fsdoc.save(f)
                
            elif any([ac,fr,at]):
                
                chartdoc.save(f)

            webbrowser.open_new(f)

        
        except IOError as e:
            
            messagebox.showinfo(title='ERROR:', message='File to be replaced is open. Could not save document.\nPlease close and try again. '+str(e))
        


# In[70]:


def openEDWARD():

    def selectFile():
        
        filetypes = (("CSV Files","*.csv"),)
    
        filename = fd.askopenfilename(
            title='Select Data File',
            initialdir='/',
            filetypes=filetypes)

        if filename:
            
            messagebox.showinfo(title='Data File:', message=filename)

            global dataFile
            dataFile = filename 
            df = pd.read_csv(dataFile).fillna(0)
            
            listdf = df[df['acct_key'] != 'CF']

            acctList = list(listdf['acct_name'].values)
            acctList.sort()
            
            account1['values'] = acctList
            account2['values'] = acctList
            account3['values'] = acctList
            
    
    def getChoices():

        xmlList = ['\"', "'", '<', '>', '&']

        escape = False

        name = nameEntry.get()

        for i in xmlList:

            if i in name:
                escape = True
                
        if not name:
            
            messagebox.showwarning("Name field incomplete:", "Please enter a company name.")
            
        elif escape:

            messagebox.showwarning("XML character detected:", "Please do not use XML characters (\", ', <, >, &).")

        else:
            
            fs = fsVar.get()
            ac = acVar.get()
            fr = frVar.get()
            at = atVar.get()

            at1 = account1.get()
            at2 = account2.get()
            at3 = account3.get()

            acclist = []

            if at1:
                acclist.append(at1)
            if at2:
                acclist.append(at2)
            if at3:
                acclist.append(at3)

            if at and not any([at1,at2,at3]):
                
                messagebox.showwarning("No trend account selected:", "If you want to produce an account trend report\nplease select at least one trend account.")
                
            else:
                
                global dataFile
                df = pd.read_csv(dataFile).fillna(0)
        
                produceDocument(fs,ac,fr,at,df, name, acclist)
    
        

    global dataFile
    
    edWindow = Tk()
    edWindow.geometry('900x850')
    edWindow.config(bg='black')
    
    edFrame = Frame(edWindow, bg='black', relief='ridge')
    edFrame.grid(row = 0)

    for i in range(12):
        edFrame.grid_rowconfigure(i, weight=1)
        
    for i in range(3):
        edFrame.grid_columnconfigure(i, weight=1)

    edlogo_img = Image.open("edbgpic.png").resize((710, 185))
    edlogo_tkimg = itk.PhotoImage(edlogo_img)
    
    edlogo = Label(edFrame, image = edlogo_tkimg, bg='gray', bd=5, relief='sunken')
    edlogo.image = edlogo_tkimg 
    edlogo.grid(row = 0, column=0, columnspan=3, pady=5)

    fileButton = Button(edFrame, text = 'Select Data File', command = selectFile, font=('OCR A Extended',15), 
                       activeforeground = 'cyan', activebackground='black', bg='gray', width = 26)
    fileButton.grid(row = 1, column = 1, pady=20, ipady=3)

    nameLabel = Label(edFrame, text = 'Please enter company name:', width = 24, font=('System',10), bg='black', fg='white')
    nameLabel.grid(row = 2, column = 1, pady = 5, ipady=2, sticky='s')
    nameEntry = Entry(edFrame, borderwidth=5, relief="ridge", width = 24, font=('System',10))
    nameEntry.grid(row = 3, column = 1, pady = 10, ipady=2)

    
    fsVar = IntVar()   
    acVar = IntVar()   
    frVar = IntVar() 
    atVar = IntVar()
      
    fsButton = Checkbutton(edFrame, text = "Balance Sheet & Income Statement",  
                          variable = fsVar, 
                          onvalue = 1, 
                          offvalue = 0, 
                          height = 2, 
                          relief='groove', 
                          font=('System',10, 'bold'),
                          activeforeground = 'cyan',
                          activebackground='black', 
                          bd=5, bg='gray', width = 30, anchor="w") 
      
    acButton = Checkbutton(edFrame, text = "Account Composition Charts", 
                          variable = acVar, 
                          onvalue = 1, 
                          offvalue = 0, 
                          height = 2, 
                          relief='groove', 
                          font=('System',10, 'bold'),
                          activeforeground = 'cyan',
                          activebackground='black', 
                          bd=5, bg='gray', width = 30, anchor="w") 
      
    frButton = Checkbutton(edFrame, text = "Financial Ratios", 
                          variable = frVar, 
                          onvalue = 1, 
                          offvalue = 0, 
                          height = 2, 
                          relief='groove', 
                          font=('System',10, 'bold'),
                          activeforeground = 'cyan',
                          activebackground='black', 
                          bd=5, bg='gray', width = 30, anchor="w")   

    atButton = Checkbutton(edFrame, text = "Account Trend Graphs", 
                          variable = atVar, 
                          onvalue = 1, 
                          offvalue = 0, 
                          height = 2, 
                          relief='groove', 
                          font=('System',10, 'bold'),
                          activeforeground = 'cyan',
                          activebackground='black',
                          bd=5, bg='gray', width = 30, anchor="w")

    fsButton.grid(row = 4, column = 1, pady=5)
    acButton.grid(row = 5, column = 1, pady=5)  
    frButton.grid(row = 6, column = 1, pady=5)
    atButton.grid(row = 7, column = 1, pady=5)

    acct1Lab = Label(edFrame, text='Trend Account 1', font=('System',8, 'bold'), fg='white', bg='black')
    acct1Lab.grid(row = 8, column = 0, pady=3, sticky ='s')
    acct2Lab = Label(edFrame, text='Trend Account 2', font=('System',8, 'bold'), fg='white', bg='black')
    acct2Lab.grid(row = 8, column = 1, pady=10, sticky ='s')
    acct3Lab = Label(edFrame, text='Trend Account 3', font=('System',8, 'bold'), fg='white', bg='black')
    acct3Lab.grid(row = 8, column = 2, pady=3, sticky ='s')
    
    account1 = ttk.Combobox(edFrame, values=[], width = 40)
    account1.grid(row = 9, column = 0, pady=5, sticky ='n')
    account2 = ttk.Combobox(edFrame, values=[], width = 40)
    account2.grid(row = 9, column = 1, pady=5, sticky ='n')
    account3 = ttk.Combobox(edFrame, values=[], width = 40)
    account3.grid(row = 9, column = 2, pady=5, sticky ='n')
    
    docButton = Button(edFrame, text = 'Produce Document', command = getChoices, font=('OCR A Extended',15), 
                       activeforeground = 'cyan', activebackground='black', bg='gray', width = 26)
    docButton.grid(row = 10, column = 1, pady=20, ipady=3)

    closeButton = Button(edFrame, text = 'Close E.D.W.A.R.D.', command = edWindow.destroy, font=('OCR A Extended',15), 
                         activeforeground = 'cyan', activebackground='black', bg='gray', width = 26)
    closeButton.grid(row = 11, column = 1, pady=15, ipady=3)

    edWindow.grid_rowconfigure(0, weight=1)    
    edWindow.grid_columnconfigure(0, weight=1)
        
    edWindow.title('E.D.W.A.R.D.')
    edWindow.mainloop()


# In[74]:


openEDWARD()

